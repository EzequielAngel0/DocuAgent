"""Extracción de texto por formato de documento.

Soporta PDF, DOCX, XLSX/XLS, CSV, MD/TXT/HTML y JSON. En PDF se inserta un
marcador `[Página N]` por página para que el chunker preserve la referencia
de página en las citas.
"""
import csv
import json
import os

from app.core.exceptions import UnsupportedFormatError


class DocumentExtractor:
    @staticmethod
    def extract_text(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return DocumentExtractor._extract_pdf(file_path)
        elif ext == ".docx":
            return DocumentExtractor._extract_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            return DocumentExtractor._extract_xlsx(file_path)
        elif ext == ".csv":
            return DocumentExtractor._extract_csv(file_path)
        elif ext in (".md", ".txt", ".html"):
            return DocumentExtractor._extract_text_file(file_path)
        elif ext == ".json":
            return DocumentExtractor._extract_json(file_path)
        else:
            raise UnsupportedFormatError(f"Extensión de archivo no soportada: {ext}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        import pypdf

        text_content = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"[Página {page_num + 1}]\n{page_text}")
        return "\n\n".join(text_content)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        import docx

        doc = docx.Document(file_path)
        text_content = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
        return "\n".join(text_content)

    @staticmethod
    def _extract_xlsx(file_path: str) -> str:
        import openpyxl

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_content = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_content.append(f"[Hoja: {sheet_name}]")
            for row in sheet.iter_rows(values_only=True):
                row_str = [str(cell).strip() for cell in row if cell is not None]
                if row_str:
                    text_content.append(", ".join(row_str))
        return "\n".join(text_content)

    @staticmethod
    def _extract_csv(file_path: str) -> str:
        text_content = []
        with open(file_path, mode="r", encoding="utf-8", errors="ignore") as f:
            for row in csv.reader(f):
                if row:
                    text_content.append(", ".join([c.strip() for c in row if c.strip()]))
        return "\n".join(text_content)

    @staticmethod
    def _extract_text_file(file_path: str) -> str:
        with open(file_path, mode="r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def _extract_json(file_path: str) -> str:
        with open(file_path, mode="r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
            return json.dumps(data, indent=2, ensure_ascii=False)
